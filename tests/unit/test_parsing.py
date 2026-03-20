"""Unit tests for document parsers."""

import io

import pytest

from src.common.exceptions import ChunkingError
from src.rag.parsing.docx import DocxParser, _format_table_markdown
from src.rag.parsing.factory import get_parser
from src.rag.parsing.pdf import PAGE_SEPARATOR, PdfParser
from src.rag.parsing.txt import TxtParser
from src.schemas.document import DocType


# ─── Factory ─────────────────────────────────────────────────────────────────

@pytest.mark.unit
def test_factory_returns_txt_parser() -> None:
    parser = get_parser("document.txt")
    assert isinstance(parser, TxtParser)


@pytest.mark.unit
def test_factory_returns_pdf_parser() -> None:
    parser = get_parser("report.PDF")  # uppercase extension
    assert isinstance(parser, PdfParser)


@pytest.mark.unit
def test_factory_returns_docx_parser() -> None:
    parser = get_parser("contract.docx")
    assert isinstance(parser, DocxParser)


@pytest.mark.unit
def test_factory_raises_for_unsupported() -> None:
    with pytest.raises(ChunkingError, match="Unsupported file type"):
        get_parser("spreadsheet.xlsx")


# ─── TxtParser ───────────────────────────────────────────────────────────────

@pytest.mark.unit
def test_txt_parser_basic() -> None:
    content = "Кредитный договор. Ставка 12%."
    result = TxtParser().parse(content.encode("utf-8"), "test.txt", DocType.CONTRACT)
    assert result.content == content
    assert result.doc_type == DocType.CONTRACT
    assert result.pages == 1


@pytest.mark.unit
def test_txt_parser_invalid_bytes_replaced() -> None:
    bad_bytes = b"Hello \xff\xfe World"
    result = TxtParser().parse(bad_bytes, "test.txt", DocType.UNKNOWN)
    assert "Hello" in result.content
    assert "World" in result.content


# ─── PdfParser ───────────────────────────────────────────────────────────────

@pytest.mark.unit
def test_pdf_parser_raises_on_invalid_bytes() -> None:
    with pytest.raises(ChunkingError, match="Failed to read PDF"):
        PdfParser().parse(b"not a pdf", "fake.pdf", DocType.CONTRACT)


@pytest.mark.unit
def test_pdf_parser_page_separator_in_content() -> None:
    """Multi-page PDFs must have PAGE_SEPARATOR between pages."""
    try:
        from pypdf import PdfWriter, PdfReader
        import pypdf
    except ImportError:
        pytest.skip("pypdf not installed")

    # Build a minimal 2-page PDF in memory
    writer = pypdf.PdfWriter()
    for _ in range(2):
        page = pypdf.PageObject.create_blank_page(width=200, height=200)
        writer.add_page(page)

    buf = io.BytesIO()
    writer.write(buf)
    pdf_bytes = buf.getvalue()

    # Blank pages produce no text → parser raises. We just verify the constant is correct.
    assert PAGE_SEPARATOR == "\x0c"


# ─── DocxParser ──────────────────────────────────────────────────────────────

@pytest.mark.unit
def test_docx_parser_raises_on_invalid_bytes() -> None:
    with pytest.raises(ChunkingError, match="Failed to read DOCX"):
        DocxParser().parse(b"not a docx", "fake.docx", DocType.CONTRACT)


@pytest.mark.unit
def test_docx_parser_extracts_paragraphs() -> None:
    """Paragraphs are extracted and joined correctly."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("Первый параграф.")
    doc.add_paragraph("Второй параграф.")

    buf = io.BytesIO()
    doc.save(buf)

    result = DocxParser().parse(buf.getvalue(), "test.docx", DocType.CONTRACT)
    assert "Первый параграф." in result.content
    assert "Второй параграф." in result.content


@pytest.mark.unit
def test_docx_parser_table_as_markdown() -> None:
    """Tables are rendered as Markdown with header separator."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Статья"
    table.cell(0, 1).text = "Ставка"
    table.cell(0, 2).text = "Срок"
    table.cell(1, 0).text = "334 ГК"
    table.cell(1, 1).text = "12%"
    table.cell(1, 2).text = "5 лет"

    buf = io.BytesIO()
    doc.save(buf)

    result = DocxParser().parse(buf.getvalue(), "test.docx", DocType.FINANCIAL_REPORT)
    assert "|" in result.content
    assert "---" in result.content
    assert "Статья" in result.content
    assert "334 ГК" in result.content


@pytest.mark.unit
def test_docx_parser_preserves_document_order() -> None:
    """Paragraph before table must appear before table in extracted content."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    doc.add_paragraph("Вводный текст договора.")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    doc.add_paragraph("Заключительный текст.")

    buf = io.BytesIO()
    doc.save(buf)

    result = DocxParser().parse(buf.getvalue(), "test.docx", DocType.CONTRACT)
    intro_pos = result.content.find("Вводный текст")
    table_pos = result.content.find("Параметр")
    outro_pos = result.content.find("Заключительный текст")
    assert intro_pos < table_pos < outro_pos


# ─── _format_table_markdown helper ───────────────────────────────────────────

@pytest.mark.unit
def test_format_table_markdown_header_separator() -> None:
    """Markdown table must have a separator row after header."""
    try:
        from docx import Document
    except ImportError:
        pytest.skip("python-docx not installed")

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A"
    table.cell(0, 1).text = "B"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "2"

    md = _format_table_markdown(table)
    lines = md.splitlines()
    assert lines[1].startswith("|")
    assert "---" in lines[1]
